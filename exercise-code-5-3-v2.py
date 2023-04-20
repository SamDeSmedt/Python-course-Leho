#!/usr/bin/python3
# python3 script_name.py --indir /path/to/input/directory --outdir /path/to/output/directory

# if no "-i, --indir" or "-o, --outdir" arguments are provided the script will use
# a default "/home/guest/Python/Files/exercise-5-3-fastqc/fastqc/" and
# "." respectively

#Import required modules
import os
import fnmatch
import re
from docx import Document
# Adding an image
# default appears at native size calculated as pixels/dpi
from docx.shared import Cm
# import argparse module
import argparse

# Functions
def get_summary(file):
    fileObj = open(file,"r")
    d = {}
    for line in fileObj.readlines():
        fields = line.split("\t")
        d[fields[1]] = fields[0]
    return d

# Argparse implementation
parser = argparse.ArgumentParser(description='Process some integers.')
parser.add_argument('-i','--indir', metavar='input_directory', type=str,
                    help='input directory for fastqc_dir')
parser.add_argument('-o','--outdir', metavar='output_directory', type=str,
                    help='output directory for fastqc_summary.docx')
args = parser.parse_args()

if args.indir:
    fastqc_dir = args.indir
else:
    fastqc_dir = "/home/guest/Python/Files/exercise-5-3-fastqc/fastqc/"

if args.outdir:
    output_dir = args.outdir
else:
    output_dir = "."

# Opening a blank document based on default template
document = Document()
# Start scanning working directory for folders to read
#print("Scanning files and folders in {}".format(fastqc_dir))
# loop over list of files and folders
fastqc_list = os.listdir(fastqc_dir)
#print(fastqc_list)
for folder in fastqc_list:
    #if dir then process
    path_folder = fastqc_dir + folder + "/"
    if (os.path.isdir(path_folder)):
        #print("Reading directory {}".format(path_folder))
        sample = folder.split("_")
        document.add_heading("Sample: {}".format(sample[0]), 0)
        document.add_heading("Summary",3)
        # Read fastqc.txt data file
        data_file = path_folder + "fastqc_data.txt"
        with open((data_file), "r") as fileObj:
                    # reading each line
            for line in fileObj.readlines():
                if re.search("^Total", line):
                    #print(line.rstrip())
                    p = document.add_paragraph(line.rstrip())
            # Read summary.txt and write in docx report
            summary_file = path_folder + "summary.txt"
            summary_dict = get_summary(summary_file)
            for k,v in summary_dict.items():
                if (k == "Basic Statistics" or k =="Per base sequence quality"):
                    summary_result = k + ": " + v
                    p = document.add_paragraph(summary_result)
            # Add image
            image = path_folder + "Images/per_base_quality.png"
            #print(image)
            document.add_picture(image, width=Cm(16))
            # Add page brake
            document.add_page_break()
# Save document
document.save('fastqc_summary.docx')