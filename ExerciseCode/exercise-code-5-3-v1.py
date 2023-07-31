#!/usr/bin/python3

import os
import fnmatch
import re
from docx import Document
# Adding an image
# default appears at native size calculated as pixels/dpi
from docx.shared import Inches, Cm

# Opening a blank document based on default template
document = Document()

# Change location to where the files are located
# Location could also be given using argparse
fastqc_dir = "/home/guest/Python/Files/exercise-5-3-fastqc/fastqc/"
os.chdir(fastqc_dir)

# Create a list with all the information that is in the folder
fastqc_list = os.listdir(fastqc_dir)

for file in fastqc_list:
    # Check if the information is a folder
    # If true, use the sampleId for the word file
    if os.path.isdir(file):
        sample_listname = file.split("_")
        # Add headers to each document
        document.add_heading("Sample: {}".format(sample_listname[0]), 0)
        document.add_heading("Summary",1)
        # Create list for folder inside each sample folder to work with
        file_dir = fastqc_dir + file + "/"
        file_list = os.listdir(file_dir)
        # Select the information necessary per folder
        for info in fnmatch.filter(file_list,"*.txt"):
            # Collect data from fastqc file
            if re.match("^fastqc", info):
                #print(info)
                with open((file_dir + info), "r") as fileObj:
                    # reading each line
                    for line in fileObj:
                        if re.search("Total Sequences",line):
                            #Inefficient as the data is splitted on the tab and concatenated on tab afterwards
                            TotalSeq = line.split("\t")
                            document.add_paragraph("{} \t {}".format(TotalSeq[0], TotalSeq[1]))
            # Collect data from the summary file
            elif re.match("summary", info):
                with open((file_dir + info), "r") as fileObj2:
                    for line in fileObj2:
                        if re.search("Basic Statistics",line) or re.search("Per base sequence quality",line):
                            TotalSeq = line.split("\t")
                            document.add_paragraph("{} \t {}".format(TotalSeq[1], TotalSeq[0]))
                            #Add empty paragraph for spacing
                    document.add_paragraph("")
        # collect the image for each sampleId
        for info in fnmatch.filter(file_list,"Images"):
            #print(info)
            # Create the directory of where all the images are stored
            images_dir = file_dir + info
            #print(images_dir)
            # Create list of all the files available in the directory
            image_list = os.listdir(images_dir)
            #print(image_list)
            for image in image_list:
                # Adjust the specified pathway for each image
                image_loc = images_dir + "/" + image
                if re.search("per_base_quality.png",image):
                    document.add_picture(image_loc, width=Cm(16))               
        # Adding a page break
        document.add_page_break()
# Save document
document.save('fastqc_summary.docx')



