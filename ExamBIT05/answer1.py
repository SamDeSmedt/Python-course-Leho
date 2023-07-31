#!/usr/bin/python3
################################################################################
# Aswer 1 from the python exam
################################################################################
import csv
import MySQLdb as my
from docx import Document
# Opening a blank document based on default template
document = Document()

# Reading data from UCSC
db = my.connect(host="genome-mysql.soe.ucsc.edu",
#db = my.connect(host="genome-euro-mysql.soe.ucsc.edu",
   user="genomep",
   passwd="password",
   db="hg19")
c = db.cursor()
# Reading the csv
with open("coordinates.csv", mode="r") as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=',')
    print("Reading region from file:")
    print("="*50)
    for row in csv_reader:
        document.add_heading("TS miRNA sites for %s" % (row[0]),2)
        split = row[0].split(":")
        pos = split[1].split("-")
        #print("row[0] = {}".format(row[0]))
        #print(split)
        #print(pos)
        print("chrom = {}\tstart = {}\tend = {}".format(split[0], pos[0], pos[1]))
        print("Getting TS miRNA sites using query:\n\
SELECT * FROM targetScanS WHERE chrom = '{}' AND chromStart > '{}' AND chromEnd < '{}'".format(split[0], pos[0], pos[1]))
        no_rows = c.execute("SELECT * FROM targetScanS WHERE chrom = '%s' AND chromStart > '%s' AND chromEnd < '%s'" % (split[0], pos[0], pos[1]))
        print("Results: {}\n".format(no_rows))
        # Fetch all results
        result = c.fetchall()
        #Iterate over each result
        for line in result:
            #print(line)
            split = line[4].split(":")
            name = split[0]
            miRNA = split[1]
            score = line[5]
            strand = line[6]
            
            #print("name = {}, miRNA = {}, positions = {}, score = {}, strand = {}".format(name, miRNA, row[0], score, strand))
            # Add table
            table = document.add_table(rows=1, cols=5)
            table.style = 'LightShading-Accent1'
            #table.style = 'LightShading-Accent1'
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'name gene'
            hdr_cells[1].text = 'name miRNA'
            hdr_cells[2].text = 'position'
            hdr_cells[3].text = 'score'
            hdr_cells[4].text = 'strand'
            # Create tuble to add data to the table
            table_data = ((name, miRNA, row[0], score, strand),)
            #print(table_data)
            
            # Add data rows
            for name, miRNA, pos, score, strand in table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = miRNA
                row_cells[2].text = pos
                row_cells[3].text = str(score)
                row_cells[4].text = strand
        # Adding a page break
        document.add_page_break()
# Save document
document.save("answer1.docx")
        